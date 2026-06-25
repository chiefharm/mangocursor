#!/usr/bin/env python3
"""Daily Mango call pipeline: auto-select, DOCX export, Telegram send."""

from __future__ import annotations

import argparse
import html as html_lib
import json
import os
import re
import subprocess
import urllib.parse
import urllib.request
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Tuple

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


PATTERNS: List[Tuple[str, str]] = [
    (r"невозможно записаться|ошибка|не да[её]т|техподдерж|сбой", "Техсбой/онлайн-запись"),
    (r"сколько стоит|стоимость|цена|дорого", "Цена/риск потери"),
    (r"перезвоню|подумаю|пока просто отменим|не смогу", "Не закрыт в запись"),
    (r"отменить запись|перезаписаться|перенести|перенос", "Перенос/отмена"),
    (r"не устраивает|позднее время|пораньше|попозже", "Неудобное время"),
    (r"недовол|жалоб|претенз|извин", "Недовольство/жалоба"),
    (r"первый раз|ранее были", "Новый клиент"),
    (r"не делаем|отказ", "Отказ в услуге"),
]


def load_dotenv(path: Path) -> None:
    if not path.exists():
        return
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def parse_transcript(html_path: Path) -> List[Tuple[str, str]]:
    src = html_path.read_text(encoding="utf-8")
    rows = re.findall(
        r'<tr>\s*<td><strong>(.*?)</strong></td>.*?<td style="width: 70%">(.*?)</td>\s*</tr>',
        src,
        flags=re.IGNORECASE | re.DOTALL,
    )
    transcript: List[Tuple[str, str]] = []
    for speaker, text_block in rows:
        cleaned = re.sub(r"<br\s*/?>", "\n", text_block, flags=re.IGNORECASE)
        cleaned = re.sub(r"<.*?>", "", cleaned, flags=re.DOTALL)
        cleaned = html_lib.unescape(cleaned)
        cleaned = "\n".join(line.strip() for line in cleaned.splitlines() if line.strip())
        if cleaned:
            transcript.append((speaker.strip(), cleaned))
    return transcript


def extract_datetime(file_name: str) -> str:
    match = re.match(r"(\d{4}-\d{2}-\d{2})__(\d{2}-\d{2}-\d{2})__", file_name)
    if not match:
        return "unknown"
    date_part, time_part = match.groups()
    return f"{date_part} {time_part.replace('-', ':')}"


def short_comment(category: str) -> str:
    c = category.lower()
    if "тех" in c:
        return "Проблема: сбой онлайн-записи. Нужно сразу переводить на ручное оформление."
    if "цена" in c:
        return "Проблема: риск потери на этапе цены. Нужен дожим до конкретного слота."
    if "не закрыт" in c or "не запис" in c:
        return "Проблема: клиент не закрыт в запись после звонка."
    if "перенос" in c or "отмен" in c:
        return "Проблема: перенос/отмена. Фиксируйте новую дату в этом же звонке."
    if "недовол" in c or "жалоб" in c:
        return "Проблема: недовольство клиента. Нужны извинение и конкретное решение."
    if "отказ" in c:
        return "Проблема: отказ в услуге. Важно предложить альтернативу."
    if "новый" in c:
        return "Фокус: первый контакт. Довести до записи и подтвердить шаг."
    return "Требуется контроль качества обработки звонка."


def classify_call(raw_html: str) -> str | None:
    for pattern, category in PATTERNS:
        if re.search(pattern, raw_html, flags=re.IGNORECASE):
            return category
    return None


def auto_select(calls_dir: Path) -> List[Dict[str, str]]:
    selected: List[Dict[str, str]] = []
    for html_path in sorted(calls_dir.glob("*.html")):
        if html_path.name.lower() == "index.html":
            continue
        raw = html_path.read_text(encoding="utf-8")
        category = classify_call(raw)
        if category:
            selected.append({"file": html_path.name, "category": category})
    return selected


def tg_send_message(token: str, chat_id: str, text: str) -> None:
    url = f"https://api.telegram.org/bot{token}/sendMessage"
    payload = urllib.parse.urlencode(
        {"chat_id": chat_id, "text": text, "disable_web_page_preview": "true"}
    ).encode("utf-8")
    req = urllib.request.Request(url, data=payload, method="POST")
    with urllib.request.urlopen(req, timeout=30) as resp:
        body = resp.read().decode("utf-8")
    if '"ok":true' not in body:
        raise RuntimeError(f"Telegram sendMessage failed: {body}")


def tg_send_document(token: str, chat_id: str, file_path: Path, caption: str) -> None:
    boundary = f"----MangoBoundary{uuid.uuid4().hex}"
    url = f"https://api.telegram.org/bot{token}/sendDocument"
    file_bytes = file_path.read_bytes()

    def part_text(name: str, value: str) -> bytes:
        return (
            f"--{boundary}\r\n"
            f'Content-Disposition: form-data; name="{name}"\r\n\r\n'
            f"{value}\r\n"
        ).encode("utf-8")

    body = bytearray()
    body.extend(part_text("chat_id", chat_id))
    body.extend(part_text("caption", caption[:1024]))
    body.extend(
        (
            f"--{boundary}\r\n"
            f'Content-Disposition: form-data; name="document"; filename="{file_path.name}"\r\n'
            "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
        ).encode("utf-8")
    )
    body.extend(file_bytes)
    body.extend(f"\r\n--{boundary}--\r\n".encode("utf-8"))

    req = urllib.request.Request(
        url,
        data=bytes(body),
        method="POST",
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
    )
    with urllib.request.urlopen(req, timeout=60) as resp:
        resp_body = resp.read().decode("utf-8")
    if '"ok":true' not in resp_body:
        raise RuntimeError(f"Telegram sendDocument failed: {resp_body}")


def write_docx(
    out_path: Path,
    idx: int,
    total: int,
    file_name: str,
    category: str,
    transcript: List[Tuple[str, str]],
) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()
    doc.add_heading(f"Звонок {idx}/{total}", level=1)
    doc.add_paragraph(f"Файл: {file_name}")
    doc.add_paragraph(f"Дата/время: {extract_datetime(file_name)}")
    doc.add_paragraph(f"Категория: {category}")
    doc.add_paragraph(f"Комментарий: {short_comment(category)}")
    doc.add_paragraph("")
    doc.add_heading("Полная расшифровка", level=2)
    for speaker, text in transcript:
        doc.add_paragraph(f"{speaker}: {text}")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def load_state(state_path: Path) -> Dict[str, str]:
    if not state_path.exists():
        return {}
    return json.loads(state_path.read_text(encoding="utf-8"))


def save_state(state_path: Path, state: Dict[str, str]) -> None:
    state_path.parent.mkdir(parents=True, exist_ok=True)
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")


def git_pull(base_dir: Path) -> None:
    if not (base_dir / ".git").exists():
        return
    subprocess.run(["git", "pull", "--ff-only"], cwd=base_dir, check=True)


def main() -> None:
    parser = argparse.ArgumentParser(description="Daily Mango call pipeline")
    parser.add_argument("--base-dir", default=".", help="Project/calls directory")
    parser.add_argument("--calls-dir", default="", help="Override calls HTML directory")
    parser.add_argument("--state-file", default="data/sent_calls.json")
    parser.add_argument("--docx-dir", default="telegram_docx")
    parser.add_argument("--dotenv", default=".env")
    parser.add_argument("--git-pull", action="store_true")
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--force", action="store_true", help="Resend even if already sent")
    args = parser.parse_args()

    base = Path(args.base_dir).resolve()
    calls_dir = Path(args.calls_dir).resolve() if args.calls_dir else base
    state_path = base / args.state_file
    docx_dir = base / args.docx_dir

    load_dotenv(base / args.dotenv)
    token = os.getenv("TELEGRAM_BOT_TOKEN", "").strip()
    chat_id = os.getenv("TELEGRAM_CHAT_ID", "").strip()

    if args.git_pull:
        git_pull(base)

    selected = auto_select(calls_dir)
    if not selected:
        print("[INFO] No problematic calls found.")
        return

    state = load_state(state_path)
    to_send: List[Dict[str, str]] = []
    for item in selected:
        file_name = item["file"]
        html_path = calls_dir / file_name
        if not html_path.exists():
            continue
        key = f"{file_name}:{html_path.stat().st_mtime_ns}"
        if not args.force and state.get(file_name) == key:
            continue
        to_send.append(item)

    if not to_send:
        print("[INFO] Nothing new to send.")
        return

    if not args.dry_run and (not token or not chat_id):
        raise SystemExit("Set TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID in .env")

    total = len(to_send)
    if not args.dry_run:
        tg_send_message(token, chat_id, f"Ежедневный отчет: найдено звонков {total}.")

    sent = 0
    for idx, item in enumerate(to_send, start=1):
        file_name = item["file"]
        category = item["category"]
        html_path = calls_dir / file_name
        transcript = parse_transcript(html_path)
        if not transcript:
            print(f"[WARN] Empty transcript: {file_name}")
            continue

        docx_name = f"call_{idx:02d}_{Path(file_name).stem}.docx"
        docx_path = docx_dir / docx_name
        write_docx(docx_path, idx, total, file_name, category, transcript)

        comment = (
            f"Звонок {idx}/{total}\n"
            f"{extract_datetime(file_name)}\n"
            f"{category}\n"
            f"{short_comment(category)}"
        )

        if args.dry_run:
            print(f"[DRY-RUN] {file_name} -> {docx_path.name}")
            continue

        tg_send_message(token, chat_id, comment)
        tg_send_document(token, chat_id, docx_path, f"DOCX {idx}/{total}")
        state[file_name] = f"{file_name}:{html_path.stat().st_mtime_ns}"
        sent += 1
        print(f"[OK] Sent {file_name}")

    if not args.dry_run:
        save_state(state_path, state)
        tg_send_message(token, chat_id, f"Готово. Отправлено DOCX: {sent}.")
        print(f"[DONE] Sent {sent} calls.")


if __name__ == "__main__":
    main()
